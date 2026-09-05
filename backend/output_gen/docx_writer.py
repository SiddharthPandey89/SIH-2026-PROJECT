"""
backend/output_gen/docx_writer.py

Local DOCX document writer for the Sovereign AI Workbench.

This module provides offline document generation for Microsoft Word .docx files
from structured content. It is a workspace-bound tool that enforces local path
security and requires no external APIs or network access.

Public API:
-----------
    write_docx(content, path, overwrite=False, **kwargs)

        Generate a .docx file from structured content.

        Args:
            content (dict): Structured document content with keys:
                - title (str, optional): Document title
                - metadata (dict, optional): Document properties
                - sections (list, optional): List of section dicts

            path (str or Path): Output file path. Must resolve within
                                PROJECT_ROOT/data/ or PROJECT_ROOT/sandbox/.

            overwrite (bool): If False (default), reject if path already exists.
                              If True, replace existing file.

            **kwargs: Reserved for future extensions.

        Returns:
            dict: Stable result shape:
                {
                    "success": bool,
                    "status": "success" | "error" | "rejected",
                    "format": "docx",
                    "path": str | None,
                    "error": str | None,
                }

Security & Constraints:
-----------------------
- Paths are validated and must remain within allowed workspace roots.
- Path traversal (../) and absolute paths outside allowed roots are rejected.
- No files are executed, evaluated, or interpreted as code.
- No environment variables or secrets are accessed.
- The implementation is completely offline and local.

Supported Content Structure:
-----------------------------
    {
        "title": "Document Title",
        "metadata": {
            "author": "...",
            "subject": "..."
        },
        "sections": [
            {
                "heading": "Section Title",
                "paragraphs": ["Text...", "More text..."],
                "bullets": ["Item 1", "Item 2"],
                "numbers": ["Step 1", "Step 2"],
                "table": {
                    "headers": ["Col 1", "Col 2"],
                    "rows": [["val1", "val2"], ["val3", "val4"]]
                },
                "page_break": True
            }
        ]
    }

All fields are optional. Missing or None fields are safely ignored.
"""

from __future__ import annotations

import logging
import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

__all__ = ["write_docx"]

logger = logging.getLogger(__name__)

PROJECT_ROOT = Path(__file__).resolve().parents[2]

ALLOWED_ROOTS: List[Path] = [
    (PROJECT_ROOT / "data").resolve(),
    (PROJECT_ROOT / "sandbox").resolve(),
]

# Constraints
_DEFAULT_MAX_FILE_SIZE = 50_000_000  # 50 MB for docx
_DEFAULT_MAX_SECTIONS = 1000
_DEFAULT_MAX_PARAGRAPHS_PER_SECTION = 1000
_DEFAULT_MAX_ROWS_PER_TABLE = 10000
_DEFAULT_MAX_COLS_PER_TABLE = 100


def _rejected_result(reason: str) -> Dict[str, Any]:
    """Build a stable rejected result dictionary."""
    return {
        "success": False,
        "status": "rejected",
        "format": "docx",
        "path": None,
        "error": reason,
    }


def _error_result(path: Optional[str], reason: str) -> Dict[str, Any]:
    """Build a stable error result dictionary."""
    return {
        "success": False,
        "status": "error",
        "format": "docx",
        "path": path,
        "error": reason,
    }


def _success_result(path: str) -> Dict[str, Any]:
    """Build a stable success result dictionary."""
    return {
        "success": True,
        "status": "success",
        "format": "docx",
        "path": path,
        "error": None,
    }


def _resolve_within_allowed(path_value: Any) -> Optional[Path]:
    """
    Resolve a path and ensure it stays within one of the allowed roots.

    Returns the resolved Path if valid, otherwise None.
    """
    if path_value is None:
        return None
    if not isinstance(path_value, (str, Path)):
        return None
    try:
        candidate = Path(path_value).expanduser().resolve()
    except (TypeError, ValueError, OSError):
        return None
    for root in ALLOWED_ROOTS:
        try:
            candidate.relative_to(root)
        except ValueError:
            continue
        return candidate
    return None


def _validate_output_extension(path: Path) -> bool:
    """
    Validate that the output path has a .docx extension.

    Returns True if valid, False otherwise.
    """
    if not path.suffix:
        return False
    return path.suffix.lower() == ".docx"


def _coerce_to_string(value: Any) -> str:
    """
    Safely convert a value to a string.

    Handles None, str, int, float, and bool.
    Returns empty string for unrecognized types.
    """
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, (int, float, bool)):
        return str(value)
    try:
        return str(value)
    except Exception:
        return ""


def _validate_content(content: Any) -> tuple[bool, Optional[str]]:
    """
    Validate the structure of the content dictionary.

    Returns (is_valid, error_message).
    """
    if content is None:
        return False, "Content cannot be None."
    if not isinstance(content, dict):
        return False, "Content must be a dictionary."

    if len(content) == 0:
        return False, "Content dictionary cannot be empty."

    # Validate title if present
    if "title" in content:
        title = content["title"]
        if title is not None and not isinstance(title, str):
            return False, "Title must be a string or None."

    # Validate metadata if present
    if "metadata" in content:
        metadata = content["metadata"]
        if metadata is not None and not isinstance(metadata, dict):
            return False, "Metadata must be a dictionary or None."

    # Validate sections if present
    if "sections" in content:
        sections = content["sections"]
        if sections is None:
            return True, None
        if not isinstance(sections, list):
            return False, "Sections must be a list or None."
        if len(sections) == 0:
            return False, "Sections list cannot be empty."
        if len(sections) > _DEFAULT_MAX_SECTIONS:
            return False, f"Too many sections (max {_DEFAULT_MAX_SECTIONS})."

        for i, section in enumerate(sections):
            if not isinstance(section, dict):
                return False, f"Section {i} must be a dictionary."
            
            # Validate section content types
            if "paragraphs" in section:
                paragraphs = section["paragraphs"]
                if paragraphs is not None and not isinstance(paragraphs, list):
                    return False, f"Section {i}: 'paragraphs' must be a list or None."
            
            if "bullets" in section:
                bullets = section["bullets"]
                if bullets is not None and not isinstance(bullets, list):
                    return False, f"Section {i}: 'bullets' must be a list or None."
            
            if "numbers" in section:
                numbers = section["numbers"]
                if numbers is not None and not isinstance(numbers, list):
                    return False, f"Section {i}: 'numbers' must be a list or None."
            
            if "table" in section:
                table = section["table"]
                if table is not None and not isinstance(table, dict):
                    return False, f"Section {i}: 'table' must be a dictionary or None."

    return True, None


def _add_title_to_doc(doc: Document, title: Optional[str]) -> None:
    """Add a title heading to the document if provided."""
    if title is None:
        return
    title_text = _coerce_to_string(title).strip()
    if not title_text:
        return
    heading = doc.add_heading(title_text, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_metadata_to_doc(doc: Document, metadata: Optional[Dict[str, Any]]) -> None:
    """Add document properties (metadata) if provided."""
    if metadata is None or not isinstance(metadata, dict):
        return

    props = doc.core_properties

    if "author" in metadata:
        author = _coerce_to_string(metadata["author"]).strip()
        if author:
            props.author = author

    if "subject" in metadata:
        subject = _coerce_to_string(metadata["subject"]).strip()
        if subject:
            props.subject = subject

    if "title" in metadata:
        title = _coerce_to_string(metadata["title"]).strip()
        if title:
            props.title = title


def _add_section_to_doc(doc: Document, section: Dict[str, Any]) -> Optional[str]:
    """
    Add a section (with heading, paragraphs, lists, tables) to the document.
    
    Returns:
        None if section added successfully
        str (error message) if validation fails and document should be rejected
    """
    if not isinstance(section, dict):
        return None

    # HARDENED: Count total paragraph/list items BEFORE generating section
    # Reject if total exceeds limit (do not silently skip items)
    total_items = 0
    
    if "paragraphs" in section:
        paragraphs = section["paragraphs"]
        if isinstance(paragraphs, list):
            for para in paragraphs:
                para_text = _coerce_to_string(para).strip()
                if para_text:
                    total_items += 1
    
    if "bullets" in section:
        bullets = section["bullets"]
        if isinstance(bullets, list):
            for item in bullets:
                item_text = _coerce_to_string(item).strip()
                if item_text:
                    total_items += 1
    
    if "numbers" in section:
        numbers = section["numbers"]
        if isinstance(numbers, list):
            for item in numbers:
                item_text = _coerce_to_string(item).strip()
                if item_text:
                    total_items += 1
    
    # Reject if total paragraph/list items exceed limit
    if total_items > _DEFAULT_MAX_PARAGRAPHS_PER_SECTION:
        return (
            f"Section has {total_items} total paragraph/list items, "
            f"which exceeds the maximum ({_DEFAULT_MAX_PARAGRAPHS_PER_SECTION}). "
            f"Section rejected. Do not silently skip items."
        )

    # All content validated, now generate section
    # Add heading if present
    if "heading" in section:
        heading = section["heading"]
        heading_text = _coerce_to_string(heading).strip()
        if heading_text:
            doc.add_heading(heading_text, level=1)

    # Add paragraphs if present
    if "paragraphs" in section:
        paragraphs = section["paragraphs"]
        if isinstance(paragraphs, list):
            for para in paragraphs:
                para_text = _coerce_to_string(para).strip()
                if para_text:
                    doc.add_paragraph(para_text)

    # Add bullet list if present
    if "bullets" in section:
        bullets = section["bullets"]
        if isinstance(bullets, list) and len(bullets) > 0:
            for item in bullets:
                item_text = _coerce_to_string(item).strip()
                if item_text:
                    doc.add_paragraph(item_text, style="List Bullet")

    # Add numbered list if present
    if "numbers" in section:
        numbers = section["numbers"]
        if isinstance(numbers, list) and len(numbers) > 0:
            for item in numbers:
                item_text = _coerce_to_string(item).strip()
                if item_text:
                    doc.add_paragraph(item_text, style="List Number")

    # Add table if present
    if "table" in section:
        table_spec = section["table"]
        table_error = _add_table_to_doc(doc, table_spec)
        if table_error:
            return table_error

    # Add page break if requested
    if section.get("page_break") is True:
        doc.add_page_break()

    return None


def _add_table_to_doc(doc: Document, table_spec: Any) -> Optional[str]:
    """
    Add a table to the document from a structured specification.
    
    Validates that all rows are list/tuple and have exactly num_cols cells.
    
    Returns:
        None if table added successfully
        str (error message) if validation fails and document should be rejected
    """
    if table_spec is None or not isinstance(table_spec, dict):
        return None

    headers = table_spec.get("headers", [])
    rows = table_spec.get("rows", [])

    if not isinstance(headers, list):
        return None

    if len(headers) == 0:
        return None

    if not isinstance(rows, list):
        return None

    num_cols = len(headers)
    num_rows = len(rows)

    # REJECT if table exceeds row limit (do not truncate)
    if num_rows > _DEFAULT_MAX_ROWS_PER_TABLE:
        return (
            f"Table has {num_rows} rows, which exceeds the maximum "
            f"({_DEFAULT_MAX_ROWS_PER_TABLE}). Table rejected."
        )

    # REJECT if table exceeds column limit (do not truncate)
    if num_cols > _DEFAULT_MAX_COLS_PER_TABLE:
        return (
            f"Table has {num_cols} columns, which exceeds the maximum "
            f"({_DEFAULT_MAX_COLS_PER_TABLE}). Table rejected."
        )

    # Validate all rows: each must be a list or tuple with exactly num_cols cells
    for row_idx, row_data in enumerate(rows):
        if not isinstance(row_data, (list, tuple)):
            return (
                f"Table row {row_idx}: not a list/tuple "
                f"(got {type(row_data).__name__}). "
                f"Each row must be a list or tuple."
            )
        if len(row_data) != num_cols:
            return (
                f"Table row {row_idx}: has {len(row_data)} cells, "
                f"expected {num_cols} (header count). "
                f"All rows must match header cell count."
            )

    # All rows are valid, create table
    if len(rows) == 0:
        # No data rows, don't create an empty table
        return None

    table = doc.add_table(rows=1 + num_rows, cols=num_cols)
    table.style = "Table Grid"

    # Add header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_text = _coerce_to_string(header)
        header_cells[i].text = header_text
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for row_idx, row_data in enumerate(rows, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx in range(num_cols):
            cell_value = row_data[col_idx]
            cell_text = _coerce_to_string(cell_value)
            row_cells[col_idx].text = cell_text

    return None


def write_docx(
    content: Any,
    path: Any,
    overwrite: bool = False,
    **kwargs: Any,
) -> Dict[str, Any]:
    """
    Generate a Microsoft Word .docx file from structured content.

    Args:
        content (dict): Structured document content.
                        See module docstring for schema.
        path (str or Path): Output file path. Must resolve within
                            PROJECT_ROOT/data/ or PROJECT_ROOT/sandbox/.
        overwrite (bool): If False (default), reject if file exists.
                          If True, replace existing file. Default: False.
        **kwargs: Reserved for future use.

    Returns:
        dict: Result dictionary with keys:
            - success (bool): True if file was written successfully.
            - status (str): "success", "error", or "rejected".
            - format (str): Always "docx".
            - path (str or None): Resolved path if successful.
            - error (str or None): Error message if unsuccessful.
    """

    # Validate and resolve path
    resolved_path = _resolve_within_allowed(path)
    if resolved_path is None:
        return _rejected_result(
            "Invalid 'path': must be inside data/ or sandbox/."
        )

    # Validate file extension
    if not _validate_output_extension(resolved_path):
        return _rejected_result(
            "Invalid file extension: must end with .docx"
        )

    # Check if file exists and overwrite policy
    if resolved_path.exists():
        if not overwrite:
            return _rejected_result(
                f"File already exists: {resolved_path}. "
                "Set overwrite=True to replace."
            )
        if resolved_path.is_dir():
            return _error_result(
                str(resolved_path),
                "Path is a directory, not a file.",
            )

    # Validate content structure
    is_valid, validation_error = _validate_content(content)
    if not is_valid:
        return _rejected_result(validation_error or "Invalid content.")

    try:
        # Create parent directories if needed
        resolved_path.parent.mkdir(parents=True, exist_ok=True)

        # Create document
        doc = Document()

        # Add title if present
        title = content.get("title")
        if title is not None:
            _add_title_to_doc(doc, title)

        # Add metadata if present
        metadata = content.get("metadata")
        if metadata is not None:
            _add_metadata_to_doc(doc, metadata)

        # Add sections if present
        sections = content.get("sections", [])
        if isinstance(sections, list):
            for section_idx, section in enumerate(sections):
                if section_idx >= _DEFAULT_MAX_SECTIONS:
                    logger.warning(
                        f"Reached maximum section limit "
                        f"({_DEFAULT_MAX_SECTIONS}), skipping remaining sections."
                    )
                    break
                if isinstance(section, dict):
                    section_error = _add_section_to_doc(doc, section)
                    if section_error:
                        # Reject document (no temp file to clean up yet)
                        return _rejected_result(
                            f"Section {section_idx} validation failed: {section_error}"
                        )

        # Use temporary file for safe atomic write
        # Write to temp file first, then rename on success
        temp_fd, temp_path = tempfile.mkstemp(
            suffix=".docx",
            dir=str(resolved_path.parent),
            prefix=".tmp_"
        )
        try:
            # Close the file descriptor
            os.close(temp_fd)
            
            # Save document to temp path
            doc.save(temp_path)

            # Verify temp file was created and has content
            temp_path_obj = Path(temp_path)
            if not temp_path_obj.exists():
                return _error_result(
                    str(resolved_path),
                    "Temporary file was not created.",
                )

            temp_file_size = temp_path_obj.stat().st_size
            if temp_file_size == 0:
                temp_path_obj.unlink(missing_ok=True)
                return _error_result(
                    str(resolved_path),
                    "Temporary file is empty (zero bytes).",
                )

            # ENFORCE FILE SIZE LIMIT: reject if temp file exceeds limit
            if temp_file_size > _DEFAULT_MAX_FILE_SIZE:
                temp_path_obj.unlink(missing_ok=True)
                return _rejected_result(
                    f"Generated file size ({temp_file_size} bytes) "
                    f"exceeds maximum limit ({_DEFAULT_MAX_FILE_SIZE} bytes)."
                )

            # Move temp file to final destination (atomic on same filesystem)
            temp_path_obj.replace(resolved_path)

        except Exception:
            # Clean up temp file on failure
            Path(temp_path).unlink(missing_ok=True)
            raise

        # Verify final file
        if not resolved_path.exists():
            return _error_result(
                str(resolved_path),
                "File was not created.",
            )

        file_size = resolved_path.stat().st_size
        if file_size == 0:
            return _error_result(
                str(resolved_path),
                "File is empty (zero bytes).",
            )

        return _success_result(str(resolved_path))

    except PermissionError as exc:
        return _error_result(
            str(resolved_path),
            f"Permission denied: {exc}",
        )

    except OSError as exc:
        return _error_result(
            str(resolved_path),
            f"File system error: {exc}",
        )

    except Exception as exc:
        logger.exception(f"Unexpected error generating DOCX: {exc}")
        return _error_result(
            str(resolved_path),
            f"Unexpected error: {type(exc).__name__}",
        )
